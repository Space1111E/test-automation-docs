import pandas as pd
from docx import Document
from docx.shared import Inches

# Lexo Excel-in
df = pd.read_excel("bug_reports.xlsx")

# Krijo dokumentin Word
doc = Document()
doc.add_heading('Bug Reports Document', 0)

for index, row in df.iterrows():
    doc.add_heading(f"{row['Bug_ID']} - {row['Title']}", level=2)

    # Informacion bazë
    doc.add_paragraph(f"🗂 Module: {row['Module']}")
    doc.add_paragraph(f"🧪 Environment: {row['Environment']}")
    doc.add_paragraph(f"👤 Reported By: {row['Reported_By']}")
    doc.add_paragraph(f"🗓 Date Reported: {row['Date_Reported']}")
    
    # Detajet e bug-ut
    doc.add_paragraph(f"\n🔁 Steps to Reproduce:\n{row['Steps_to_Reproduce']}")
    doc.add_paragraph(f"❌ Actual Result:\n{row['Actual_Result']}")
    doc.add_paragraph(f"✅ Expected Result:\n{row['Expected_Result']}")
    
    # Vlerësimi
    doc.add_paragraph(f"\n🚨 Severity: {row['Severity']}")
    doc.add_paragraph(f"🎯 Priority: {row['Priority']}")
    doc.add_paragraph(f"📌 Status: {row['Status']}")
    doc.add_paragraph(f"📝 Notes: {row['Notes']}")
    
    doc.add_paragraph("-" * 40)

# Ruaje dokumentin
doc.save("Generated_Bug_Reports.docx")
print("✅ Bug reports u gjeneruan me sukses!")
